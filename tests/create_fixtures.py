import docx
import os

def create_fixtures(output_dir):
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    # 1. has_display_math.docx
    doc = docx.Document()
    doc.add_paragraph("Some text before")
    doc.add_paragraph("$$A=1$$")
    doc.add_paragraph("Some text after")
    doc.save(os.path.join(output_dir, "has_display_math.docx"))

    # 2. escaped_dollar.docx
    doc = docx.Document()
    doc.add_paragraph(r"Here is an escaped dollar: \$ $100")
    doc.add_paragraph(r"And another: \$$ not display math")
    doc.save(os.path.join(output_dir, "escaped_dollar.docx"))

    # 3. inline_math_only.docx
    doc = docx.Document()
    doc.add_paragraph("This is inline math $x+y$ which is fine.")
    doc.save(os.path.join(output_dir, "inline_math_only.docx"))

    # 4. split_runs_display.docx
    doc = docx.Document()
    p = doc.add_paragraph()
    r1 = p.add_run("Here is broken display math ")
    r2 = p.add_run("$")
    r3 = p.add_run("$") # consecutive runs form $$
    r4 = p.add_run("B=2")
    r5 = p.add_run("$")
    r6 = p.add_run("$")
    doc.save(os.path.join(output_dir, "split_runs_display.docx"))

    # 5. no_math.docx
    doc = docx.Document()
    doc.add_paragraph("Just plain text file.")
    doc.save(os.path.join(output_dir, "no_math.docx"))
    
    # 6. invalid file (not doing it here, will just touch a file)
    with open(os.path.join(output_dir, "not_a_docx.txt"), "w") as f:
        f.write("I am not a docx")

if __name__ == "__main__":
    create_fixtures("d:/thinksolv/tast1_fr/fixtures")
